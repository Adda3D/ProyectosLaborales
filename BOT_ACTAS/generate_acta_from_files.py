# generate_acta_from_files.py
import os, argparse, requests
from pathlib import Path
import docx2txt
from docx import Document

def read_text_or_docx(p: Path) -> str:
    if p.suffix.lower()==".docx": return docx2txt.process(str(p)) or ""
    return p.read_text(encoding="utf-8", errors="replace")

def md_to_docx(md: str, out: Path):
    doc = Document()
    for raw in md.splitlines():
        line = raw.rstrip()
        if line.startswith("# "):   doc.add_heading(line[2:].strip(), level=1)
        elif line.startswith("## "): doc.add_heading(line[3:].strip(), level=2)
        elif line.startswith("### "):doc.add_heading(line[4:].strip(), level=3)
        elif line.startswith("- "):
            p = doc.add_paragraph(); p.style="List Bullet"; p.add_run(line[2:].strip())
        else: doc.add_paragraph(line)
    doc.save(str(out))

def chat(messages, api_key, model="deepseek-chat"):
    url="https://api.deepseek.com/v1/chat/completions"
    r=requests.post(url,
        headers={"Authorization":f"Bearer {api_key}","Content-Type":"application/json"},
        json={"model":model,"messages":messages,"temperature":0.2,"max_tokens":6000},
        timeout=180)
    r.raise_for_status()
    return r.json()["choices"][0]["message"]["content"].strip()

if __name__=="__main__":
    ap=argparse.ArgumentParser(description="Genera acta (MD/DOCX) a partir de .spk.txt + convocatoria/prompt")
    ap.add_argument("--spk", required=True, help="Transcripción diarizada (.spk.txt)")
    ap.add_argument("--convocatoria", required=True, help="Convocatoria .docx")
    ap.add_argument("--prompt", required=True, help="Prompt base .txt")
    ap.add_argument("--acta-ejemplo", default="", help="Acta de ejemplo .docx (opcional)")
    ap.add_argument("--out", required=True, help="Prefijo de salida (sin extensión)")
    ap.add_argument("--comite", required=True)
    ap.add_argument("--acta", default="[NÚMERO]")
    ap.add_argument("--sesion", default="Presencial")
    ap.add_argument("--fecha", required=True)
    ap.add_argument("--inicio", default="")
    ap.add_argument("--cierre", default="No consta")
    ap.add_argument("--lugar", required=True)
    ap.add_argument("--api-key", default=os.getenv("DEEPSEEK_API_KEY",""))
    args=ap.parse_args()

    if not args.api_key: raise SystemExit("Falta --api-key o variable DEEPSEEK_API_KEY.")

    conv = read_text_or_docx(Path(args.convocatoria))
    ejemplo = read_text_or_docx(Path(args.acta_ejemplo)) if args.acta_ejemplo else "[No provisto]"
    prompt = Path(args.prompt).read_text(encoding="utf-8")
    spk = Path(args.spk).read_text(encoding="utf-8")

    system = (
        "Eres un redactor experto de actas universitarias. "
        "Redacta siempre en tercera persona, sobria y clara. "
        "Sigue la estructura y el estilo del documento de ejemplo si está disponible. "
        "No inventes datos: usa [DATO NO CONSTA] cuando falte. "
        "Entrega: Acta en Markdown + Resumen ejecutivo + tablas de decisiones/votaciones + compromisos.\n\n"
        "Reglas especiales sobre intervenciones:\n"
        "- La transcripción usa etiquetas como 'Persona 1', 'Persona 2', etc.\n"
        "- En el apartado 'Desarrollo de la sesión', cada intervención debe aparecer en este formato:\n"
        "  *Persona 1:* expuso tal cosa.  *Persona 2:* opinó sobre lo anterior.\n"
        "- Si se identifica el nombre o rol de la persona en la transcripción, reemplaza 'Persona X' por ese nombre/rol.\n"
        "- Si no se puede identificar, conserva 'Persona X' o usa 'Interviniente [X]'.\n\n"
        "--- PROMPT BASE ---\n" + prompt
    )

    meta = {
        "COMITÉ": args.comite,
        "ACTA": args.acta,
        "SESIÓN": args.sesion,
        "FECHA": args.fecha,
        "HORA INICIO": args.inicio,
        "HORA CIERRE": args.cierre,
        "LUGAR": args.lugar,
    }

    user = (
        "## Metadatos\n" + "\n".join(f"- {k}: {v}" for k, v in meta.items()) +
        "\n\n## Convocatoria (texto)\n" + conv[:120000] +
        "\n\n## Documento de Ejemplo (texto)\n" + ejemplo[:120000] +
        "\n\n## Transcripción diarizada\n" + spk[:350000] +
        "\n\n### Tarea\n"
        "Elabora el acta completa en Markdown con el orden del día de la convocatoria, "
        "ajustando el desarrollo a lo indicado en la transcripción. "
        "Incluye sección de votaciones/decisiones (en tabla) y compromisos (si los hay)."
    )

    md = chat([{"role":"system","content":system},{"role":"user","content":user}], api_key=args.api_key)

    out_md = Path(args.out).with_suffix(".md"); out_md.write_text(md, encoding="utf-8")
    out_docx = Path(args.out).with_suffix(".docx"); md_to_docx(md, out_docx)
    print("OK ->", out_md, "|", out_docx)
