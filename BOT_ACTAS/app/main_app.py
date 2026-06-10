import os, sys, threading, subprocess, requests, datetime as dt
import docx2txt
from docx import Document
from pathlib import Path
import tkinter as tk
from tkinter import ttk, filedialog, messagebox

from app.utils_paths import resource_path
from scripts.transcribe_whisper_local import main as whisper_main

APP_TITLE = "Generador de Actas – FDCPS"
DEFAULT_SPEAKERS = 2

# ===================== Utilidades =====================
def which_ffmpeg() -> str | None:
    """
    Busca ffmpeg: 1) dentro del bundle (app/ffmpeg/ffmpeg.exe), 2) PATH, 3) C:\ffmpeg\...\bin.
    """
    # 1) embebido
    cand = resource_path("app/ffmpeg/ffmpeg.exe")
    if cand.exists():
        return str(cand)

    # 2) PATH
    from shutil import which as _which
    p = _which("ffmpeg")
    if p:
        return p

    # 3) C:\ffmpeg
    try:
        base = Path("C:/ffmpeg")
        if base.exists():
            latest = sorted(base.glob("*/bin/ffmpeg.exe"),
                            key=lambda q: q.stat().st_mtime, reverse=True)
            if latest:
                return str(latest[0])
    except Exception:
        pass
    return None

def run_cmd(cmd: list[str], cwd: str | None = None) -> int:
    """Ejecuta comando y muestra la salida en la consola (logs)."""
    print("[run]", " ".join(cmd))
    proc = subprocess.Popen(
        cmd, cwd=cwd,
        stdout=subprocess.PIPE, stderr=subprocess.STDOUT,
        text=True, encoding="utf-8", errors="replace"
    )
    while True:
        line = proc.stdout.readline()
        if not line and proc.poll() is not None:
            break
        if line:
            print(line.rstrip())
    rc = proc.returncode
    print(f"[run] returncode={rc}")
    return rc

def safe_name(s: str) -> str:
    keep = "-_.() "
    return "".join(ch for ch in s if ch.isalnum() or ch in keep).strip().replace(" ", "_")

def read_text_or_docx(path: Path) -> str:
    p = Path(path)
    if not p.exists():
        return ""
    if p.suffix.lower() == ".docx":
        return docx2txt.process(str(p)) or ""
    return p.read_text(encoding="utf-8", errors="replace")

def load_prompt_file(path: Path) -> str:
    try:
        return Path(path).read_text(encoding="utf-8")
    except Exception:
        return (
            "Eres un redactor experto en actas universitarias. Redacta en tercera persona. "
            "Usa el orden del día como esqueleto y ajusta al desarrollo real según la transcripción. "
            "No inventes: si falta algo usa [DATO NO CONSTA]. Incluye asistentes, orden del día, "
            "desarrollo, decisiones/votaciones, compromisos y resumen ejecutivo. Devuelve en Markdown."
        )

def call_deepseek_chat(messages: list[dict], model: str | None = None, api_key: str | None = None) -> str:
    base = "https://api.deepseek.com/v1"
    url = f"{base}/chat/completions"
    model = model or os.getenv("DEEPSEEK_MODEL", "deepseek-chat")
    api_key = api_key or os.getenv("DEEPSEEK_API_KEY", "")
    if not api_key:
        raise RuntimeError("Falta DEEPSEEK_API_KEY en variables de entorno (o complétalo en la app).")
    headers = {"Authorization": f"Bearer {api_key}", "Content-Type": "application/json"}
    payload = {"model": model, "messages": messages, "temperature": 0.2, "max_tokens": 6000}
    print("[deepseek] POST /chat/completions")
    r = requests.post(url, headers=headers, json=payload, timeout=180)
    r.raise_for_status()
    return r.json()["choices"][0]["message"]["content"].strip()

def md_to_docx(md_text: str, out_docx: Path):
    doc = Document()
    for raw in md_text.splitlines():
        line = raw.rstrip()
        if line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=1)
        elif line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=2)
        elif line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=3)
        elif line.startswith("- "):
            p = doc.add_paragraph()
            p.style = "List Bullet"
            p.add_run(line[2:].strip())
        else:
            doc.add_paragraph(line)
    doc.save(str(out_docx))

def set_container_state(container: tk.Misc, enabled: bool):
    for child in container.winfo_children():
        if isinstance(child, (ttk.Frame, ttk.LabelFrame, tk.Frame, tk.LabelFrame)):
            set_container_state(child, enabled)
            continue
        try:
            if hasattr(child, "state"):
                child.state(["!disabled"] if enabled else ["disabled"])
                continue
        except Exception:
            pass
        try:
            child.configure(state=("normal" if enabled else "disabled"))
        except Exception:
            pass

# ===================== GUI =====================
class App(tk.Tk):
    def __init__(self):
        super().__init__()
        self.title(APP_TITLE)
        self.geometry("920x760")
        self.minsize(900, 720)

        self.use_existing_var = tk.BooleanVar(value=False)
        self.existing_txt_var = tk.StringVar(value="")

        self.file_var = tk.StringVar()
        self.outdir_var = tk.StringVar(value=str(Path.cwd() / "salidas"))
        self.name_var = tk.StringVar(value="")
        self.spk_var = tk.IntVar(value=DEFAULT_SPEAKERS)
        self.status_var = tk.StringVar(value="Listo.")

        self.committee_var = tk.StringVar(value="")
        self.acta_num_var = tk.StringVar(value="")
        self.session_var = tk.StringVar(value="Presencial")
        self.date_var = tk.StringVar(value=dt.date.today().strftime("%d/%m/%Y"))
        self.start_var = tk.StringVar(value="")
        self.end_var = tk.StringVar(value="")
        self.place_var = tk.StringVar(value="")

        self.convocatoria_var = tk.StringVar(value="")
        self.acta_ejemplo_var = tk.StringVar(value="")
        # usa prompt embebido por defecto
        self.prompt_var = tk.StringVar(value=str(resource_path("app/promp.txt")))

        self.api_key_var = tk.StringVar(value=os.getenv("DEEPSEEK_API_KEY", ""))

        self.toast = None
        self._toast_ok = False

        self._build_ui()

    def _build_ui(self):
        title = ttk.Label(self, text=APP_TITLE, font=("Segoe UI", 18, "bold"))
        title.pack(anchor="w", padx=12, pady=(12, 8))

        sbar = ttk.Frame(self); sbar.pack(fill="x", padx=12, pady=(0, 6))
        ttk.Checkbutton(
            sbar,
            variable=self.use_existing_var,
            text="Ya tengo la transcripción diarizada (.spk.txt) y solo quiero generar el Acta"
        ).pack(anchor="w")
        self.use_existing_var.trace_add("write", lambda *_: self._toggle_flow())

        self.flow_a = ttk.LabelFrame(self, text="Flujo A: crear transcripción desde audio")
        self.flow_a.pack(fill="x", padx=12, pady=6)
        r1 = ttk.Frame(self.flow_a); r1.pack(fill="x", pady=4)
        ttk.Label(r1, text="Archivo de audio (.wav):").pack(side="left", padx=(0, 10))
        ttk.Entry(r1, textvariable=self.file_var).pack(side="left", fill="x", expand=True)
        ttk.Button(r1, text="Buscar…", command=self._pick_wav).pack(side="left", padx=6)

        r2 = ttk.Frame(self.flow_a); r2.pack(fill="x", pady=4)
        ttk.Label(r2, text="Carpeta de salida:").pack(side="left", padx=(0, 10))
        ttk.Entry(r2, textvariable=self.outdir_var).pack(side="left", fill="x", expand=True)
        ttk.Button(r2, text="Elegir…", command=self._pick_outdir).pack(side="left", padx=6)

        r3 = ttk.Frame(self.flow_a); r3.pack(fill="x", pady=4)
        ttk.Label(r3, text="Nombre del resultado:").pack(side="left", padx=(0, 10))
        ttk.Entry(r3, textvariable=self.name_var).pack(side="left", fill="x", expand=True)

        r4 = ttk.Frame(self.flow_a); r4.pack(fill="x", pady=4)
        ttk.Label(r4, text="N.º de personas (speakers):").pack(side="left", padx=(0, 10))
        ttk.Spinbox(r4, from_=1, to=16, textvariable=self.spk_var, width=6).pack(side="left")

        self.flow_b = ttk.LabelFrame(self, text="Flujo B: usar transcripción existente (.spk.txt)")
        self.flow_b.pack(fill="x", padx=12, pady=6)
        rb = ttk.Frame(self.flow_b); rb.pack(fill="x", pady=4)
        ttk.Label(rb, text="Archivo .spk.txt:").pack(side="left", padx=(0, 10))
        ttk.Entry(rb, textvariable=self.existing_txt_var).pack(side="left", fill="x", expand=True)
        ttk.Button(rb, text="Elegir…", command=self._pick_spk).pack(side="left", padx=6)

        btns = ttk.Frame(self); btns.pack(fill="x", padx=12, pady=(8, 8))
        self.btn_go = ttk.Button(btns, text="Iniciar", command=self._start)
        self.btn_go.pack(side="left")
        ttk.Button(btns, text="Minimizar", command=self.iconify).pack(side="left", padx=8)
        ttk.Button(btns, text="Salir", command=self.destroy).pack(side="right")

        ttk.Separator(self, orient="horizontal").pack(fill="x", padx=12, pady=8)

        ttk.Label(self, text="Metadatos del Acta", font=("Segoe UI", 12, "bold")).pack(anchor="w", padx=12)
        grid = ttk.Frame(self); grid.pack(fill="x", padx=12, pady=(4, 8))

        def add_row(lbl, var, width=34):
            r = ttk.Frame(grid); r.pack(fill="x", pady=2)
            ttk.Label(r, text=lbl, width=28, anchor="w").pack(side="left")
            ttk.Entry(r, textvariable=var, width=width).pack(side="left", fill="x", expand=True)

        add_row("COMITÉ:", self.committee_var)
        add_row("ACTA (N.º):", self.acta_num_var)
        add_row("SESIÓN (Presencial/Mixta/Virtual):", self.session_var)
        add_row("FECHA (dd/mm/aaaa):", self.date_var)
        add_row("HORA INICIO (hh:mm):", self.start_var)
        add_row("HORA CIERRE (hh:mm o 'No consta'):", self.end_var)
        add_row("LUGAR:", self.place_var)

        ttk.Separator(self, orient="horizontal").pack(fill="x", padx=12, pady=8)

        ttk.Label(self, text="Documentos guía", font=("Segoe UI", 12, "bold")).pack(anchor="w", padx=12)
        g2 = ttk.Frame(self); g2.pack(fill="x", padx=12, pady=(4, 8))

        def file_row(lbl, var, picker):
            r = ttk.Frame(g2); r.pack(fill="x", pady=3)
            ttk.Label(r, text=lbl, width=28, anchor="w").pack(side="left")
            ttk.Entry(r, textvariable=var).pack(side="left", fill="x", expand=True)
            ttk.Button(r, text="Elegir…", command=picker).pack(side="left", padx=6)

        file_row("Convocatoria (DOCX):", self.convocatoria_var, self._pick_convocatoria)
        file_row("Acta de ejemplo (DOCX):", self.acta_ejemplo_var, self._pick_acta_ejemplo)
        file_row("Prompt base (TXT):", self.prompt_var, self._pick_prompt)

        ttk.Separator(self, orient="horizontal").pack(fill="x", padx=12, pady=8)

        ak = ttk.Frame(self); ak.pack(fill="x", padx=12, pady=(0, 6))
        ttk.Label(ak, text="DeepSeek API Key (Si se vence:)", anchor="w").pack(side="left")
        ttk.Entry(ak, textvariable=self.api_key_var, show="•", width=60).pack(side="left", padx=8)

        self.prog = ttk.Progressbar(self, mode="indeterminate"); self.prog.pack(fill="x", padx=12)
        ttk.Label(self, textvariable=self.status_var).pack(anchor="w", padx=12, pady=(6, 6))
        ttk.Label(self, text="Software realizado por Oficina de Comunicaciones y Tecnología FDCPS - 2025 - Puedes minimizar, pero no cierres la app; te avisaremos cuando esté listo.",
                  foreground="#444").pack(anchor="w", padx=12)

        self._toggle_flow()

    # ===================== Pickers =====================
    def _pick_spk(self):
        p = filedialog.askopenfilename(
            title="Seleccionar transcripción (.spk.txt)",
            filetypes=[("Transcripción diarizada", "*.spk.txt"), ("Texto", "*.txt")]
        )
        if p:
            self.existing_txt_var.set(p)

    def _pick_convocatoria(self):
        p = filedialog.askopenfilename(title="Seleccionar Convocatoria (DOCX)", filetypes=[("Word", "*.docx")])
        if p:
            self.convocatoria_var.set(p)

    def _pick_acta_ejemplo(self):
        p = filedialog.askopenfilename(title="Seleccionar Acta de ejemplo (DOCX)", filetypes=[("Word", "*.docx")])
        if p:
            self.acta_ejemplo_var.set(p)

    def _pick_prompt(self):
        p = filedialog.askopenfilename(title="Seleccionar prompt base (TXT)", filetypes=[("Texto", "*.txt")])
        if p:
            self.prompt_var.set(p)

    def _pick_wav(self):
        p = filedialog.askopenfilename(title="Seleccionar WAV", filetypes=[("WAV", "*.wav")])
        if p:
            self.file_var.set(p)
            if not self.name_var.get():
                self.name_var.set(safe_name(Path(p).stem))

    def _pick_outdir(self):
        d = filedialog.askdirectory(title="Elegir carpeta de salida")
        if d:
            self.outdir_var.set(d)

    # ===================== Toggle de flujos =====================
    def _toggle_flow(self):
        use_existing = self.use_existing_var.get()
        set_container_state(self.flow_a, enabled=not use_existing)
        set_container_state(self.flow_b, enabled=use_existing)

    # ===================== Validación =====================
    def _validate(self) -> tuple[bool, str]:
        outdir = Path(self.outdir_var.get() or ".")
        name = safe_name(self.name_var.get() or "resultado")
        outdir.mkdir(parents=True, exist_ok=True)

        if self.use_existing_var.get():
            spk = Path(self.existing_txt_var.get())
            if not spk.is_file():
                return False, "Selecciona el archivo .spk.txt existente."
            if not spk.name.endswith(".spk.txt"):
                return False, "El archivo debe terminar en .spk.txt."
        else:
            src = Path(self.file_var.get())
            if not src.is_file() or src.suffix.lower() != ".wav":
                return False, "Selecciona un archivo WAV válido."
            if int(self.spk_var.get() or 0) < 1:
                return False, "El número de personas debe ser ≥ 1."

        conv = Path(self.convocatoria_var.get())
        if not (conv.is_file() and conv.suffix.lower() == ".docx"):
            return False, "Selecciona la convocatoria (.docx)."

        pmt = Path(self.prompt_var.get())
        if not pmt.is_file():
            return False, "Selecciona el prompt base (promp.txt)."

        if not self.committee_var.get().strip(): return False, "Completa el campo COMITÉ."
        if not self.session_var.get().strip():   return False, "Completa el campo SESIÓN."
        if not self.date_var.get().strip():      return False, "Completa la FECHA."
        if not self.place_var.get().strip():     return False, "Completa el LUGAR."

        ak = self.api_key_var.get().strip()
        if ak:
            os.environ["DEEPSEEK_API_KEY"] = ak

        return True, ""

    # ===================== Estado =====================
    def _set_busy(self, busy: bool, msg: str | None = None):
        if busy:
            self.btn_go.state(["disabled"])
            self.prog.start(12)
        else:
            self.btn_go.state(["!disabled"])
            self.prog.stop()
        if msg:
            self.status_var.set(msg)

    # ===================== Ejecución =====================
    def _start(self):
        ok, err = self._validate()
        if not ok:
            messagebox.showwarning(APP_TITLE, err)
            return
        self._set_busy(True, "Procesando… puedes minimizar, te avisaremos al terminar.")
        threading.Thread(target=self._run_pipeline, daemon=True).start()

    def _run_pipeline(self):
        try:
            outdir = Path(self.outdir_var.get()).resolve()
            name = safe_name(self.name_var.get() or "resultado")

            if self.use_existing_var.get():
                spk_txt = Path(self.existing_txt_var.get()).resolve()
                final_txt = outdir / f"{name}.spk.txt"
                if spk_txt != final_txt:
                    final_txt.write_text(spk_txt.read_text(encoding="utf-8"), encoding="utf-8")
            else:
                src = Path(self.file_var.get()).resolve()
                speakers = int(self.spk_var.get() or DEFAULT_SPEAKERS)

                self._update("Convirtiendo a 16 kHz mono con FFmpeg…")
                ff = which_ffmpeg()
                if not ff:
                    raise RuntimeError("No se encontró ffmpeg (embebido ni en PATH).")
                wav16 = outdir / f"{name}_16k.wav"
                rc = run_cmd([ff, "-y", "-i", str(src), "-ac", "1", "-ar", "16000", "-c:a", "pcm_s16le", str(wav16)])
                if rc != 0:
                    raise RuntimeError("Falló la conversión con ffmpeg.")

                self._update("Transcribiendo con Whisper (local)…")
                os.environ["KMP_DUPLICATE_LIB_OK"] = "TRUE"
                whisper_json = outdir / f"{name}_16k.whisper.json"
                rc = whisper_main(file=str(wav16), model="medium", lang="es", out_json=str(whisper_json))
                if rc != 0 or not whisper_json.exists():
                    raise RuntimeError("Falló la transcripción con Whisper.")
                
                try:
                    from scripts.diarize_speechbrain_merge import main as diarize_main
                except Exception as e:
                    raise RuntimeError(
                        "No se pudo cargar el módulo de diarización (speechbrain/sklearn). "
                        "Revisa el bundle o las dependencias.\n\nDetalle: " + str(e)
                    )

                self._update(f"Diarizando y fusionando (speakers={speakers})…")
                spk_txt = outdir / f"{name}.spk.txt"
                rc = diarize_main(wav=str(wav16), json_path=str(whisper_json), speakers=speakers, out_spk=str(spk_txt))
                if rc != 0 or not spk_txt.exists():
                    raise RuntimeError("Falló la diarización/fusión.")


                final_txt = outdir / f"{name}.spk.txt"
                if spk_txt.resolve() != final_txt.resolve():
                    final_txt.write_text(spk_txt.read_text(encoding="utf-8"), encoding="utf-8")

            self._update("Generando acta con IA…")
            conv_text = read_text_or_docx(Path(self.convocatoria_var.get()))
            acta_ej_text = read_text_or_docx(Path(self.acta_ejemplo_var.get())) if self.acta_ejemplo_var.get().strip() else ""
            prompt_base = load_prompt_file(Path(self.prompt_var.get()))
            transcript_text = (outdir / f"{name}.spk.txt").read_text(encoding="utf-8", errors="replace")

            system_msg = (
                "Eres un redactor experto de actas universitarias. "
                "Redacta siempre en tercera persona, sobria y clara. "
                "Sigue la estructura y el estilo del documento de ejemplo si está disponible. "
                "No inventes datos: usa [DATO NO CONSTA] cuando falte. "
                "Entrega: Acta en Markdown + Resumen ejecutivo + tablas de decisiones/votaciones + compromisos.\n\n"
                "Reglas sobre intervenciones:\n"
                "- La transcripción usa etiquetas como 'Persona 1', 'Persona 2', etc.\n"
                "- En 'Desarrollo de la sesión', refleja las intervenciones con el formato: "
                "*Persona 1:* expuso tal cosa. *Persona 2:* opinó sobre lo anterior.\n"
                "- Si se identifica el nombre o rol en la transcripción, reemplaza 'Persona X' por ese nombre/rol.\n"
                "- Si no, conserva 'Persona X' o usa 'Interviniente [X]'.\n\n"
                "--- PROMPT BASE ---\n" + (prompt_base or "")
            )

            meta = {
                "COMITÉ": self.committee_var.get().strip(),
                "ACTA": self.acta_num_var.get().strip() or "[NÚMERO]",
                "SESIÓN": self.session_var.get().strip(),
                "FECHA": self.date_var.get().strip(),
                "HORA INICIO": self.start_var.get().strip(),
                "HORA CIERRE": self.end_var.get().strip() or "No consta",
                "LUGAR": self.place_var.get().strip(),
            }
            user_msg = (
                "## Metadatos\n" + "\n".join(f"- {k}: {v}" for k, v in meta.items()) +
                "\n\n## Convocatoria (texto)\n" + conv_text[:120000] +
                "\n\n## Documento de Ejemplo (texto)\n" + (acta_ej_text[:120000] if acta_ej_text else "[No provisto]") +
                "\n\n## Transcripción diarizada\n" + transcript_text[:350000] +
                "\n\n### Tarea\nElabora el acta completa (Markdown) con el orden del día de la convocatoria, "
                "ajustando el desarrollo a lo indicado por la transcripción. Incluye votaciones/decisiones y compromisos."
            )

            md = call_deepseek_chat(
                messages=[{"role": "system", "content": system_msg},
                          {"role": "user", "content": user_msg}],
                api_key=self.api_key_var.get().strip() or None
            )

            acta_md = outdir / f"{name}.acta.md"
            acta_md.write_text(md, encoding="utf-8")
            acta_docx = outdir / f"{name}.acta.docx"
            md_to_docx(md, acta_docx)

            self._finish_ok(acta_docx)

        except Exception as e:
            self._finish_err(str(e))

    # ===================== Helpers GUI =====================
    def _update(self, msg: str):
        print("[status]", msg)
        self.after(0, lambda: self.status_var.set(msg))

    def _finish_ok(self, p: Path):
        def _d():
            self._set_busy(False, f"¡Listo! Resultado: {p}")
            try:
                if self._toast_ok:
                    self.toast.show_toast(APP_TITLE, f"Proceso completo:\n{p}", duration=6, threaded=False)
            except Exception as e:
                print("[toast] fallo:", e)
            messagebox.showinfo(APP_TITLE, f"Proceso completo.\n\n{p}")
        self.after(0, _d)

    def _finish_err(self, err: str):
        def _d():
            self._set_busy(False, f"Error: {err}")
            try:
                if self._toast_ok:
                    self.toast.show_toast(APP_TITLE, f"Error: {err}", duration=8, threaded=False)
            except Exception as e:
                print("[toast] fallo:", e)
            messagebox.showerror(APP_TITLE, f"Ocurrió un error:\n\n{err}")
        self.after(0, _d)


# ===================== Main =====================
if __name__ == "__main__":
    App().mainloop()
