# run_all.py
import os, sys, json, subprocess, argparse, datetime as dt
from pathlib import Path
import requests
import docx2txt
from docx import Document

# ---------- Utilidades ----------
def which_ffmpeg() -> str | None:
    from shutil import which
    p = which("ffmpeg")
    if p: return p
    # fallback típico de Windows host (no aplica en Docker, pero no estorba)
    base = Path("C:/ffmpeg")
    if base.exists():
        latest = sorted(base.glob("*/bin/ffmpeg.exe"), key=lambda q: q.stat().st_mtime, reverse=True)
        if latest: return str(latest[0])
    return None

def run_cmd(cmd: list[str]) -> int:
    proc = subprocess.Popen(cmd, stdout=subprocess.PIPE, stderr=subprocess.STDOUT, text=True, encoding="utf-8", errors="replace")
    for line in proc.stdout:
        print(line.rstrip())
    return proc.wait()

def read_text_or_docx(path: Path) -> str:
    if not path or not path.exists(): return ""
    if path.suffix.lower() == ".docx":
        return docx2txt.process(str(path)) or ""
    return path.read_text(encoding="utf-8", errors="replace")

def md_to_docx(md_text: str, out_docx: Path):
    doc = Document()
    for raw in md_text.splitlines():
        line = raw.rstrip()
        if line.startswith("# "):   doc.add_heading(line[2:].strip(), level=1)
        elif line.startswith("## "): doc.add_heading(line[3:].strip(), level=2)
        elif line.startswith("### "):doc.add_heading(line[4:].strip(), level=3)
        elif line.startswith("- "):
            p = doc.add_paragraph(); p.style = "List Bullet"; p.add_run(line[2:].strip())
        else:
            doc.add_paragraph(line)
    doc.save(str(out_docx))

def call_deepseek_chat(messages, api_key: str, model: str = "deepseek-chat", temperature: float = 0.2, max_tokens: int = 6000) -> str:
    url = "https://api.deepseek.com/v1/chat/completions"
    headers = {"Authorization": f"Bearer {api_key}", "Content-Type": "application/json"}
    payload = {"model": model, "messages": messages, "temperature": temperature, "max_tokens": max_tokens}
    r = requests.post(url, headers=headers, json=payload, timeout=180)
    r.raise_for_status()
    return r.json()["choices"][0]["message"]["content"].strip()

# ---------- Pipeline ----------
def main():
    ap = argparse.ArgumentParser(description="Transcribe + Diarize + Generate Minutes (Acta)")
    gsrc = ap.add_mutually_exclusive_group(required=True)
    gsrc.add_argument("--wav", type=str, help="Ruta al WAV de entrada")
    gsrc.add_argument("--spk", type=str, help="Ruta a la transcripción diarizada existente (.spk.txt)")

    ap.add_argument("--outdir", type=str, required=True, help="Carpeta de salida")
    ap.add_argument("--name", type=str, required=True, help="Nombre base de salida (sin extensión)")
    ap.add_argument("--speakers", type=int, default=2, help="Nº de personas (si parte de WAV)")

    ap.add_argument("--convocatoria", type=str, required=True, help="DOCX de convocatoria")
    ap.add_argument("--acta-ejemplo", type=str, default="", help="DOCX de ejemplo (opcional)")
    ap.add_argument("--prompt", type=str, required=True, help="prompt base (TXT)")
    ap.add_argument("--api-key", type=str, default=os.getenv("DEEPSEEK_API_KEY", ""), help="DeepSeek API key")

    # metadatos
    ap.add_argument("--comite", type=str, required=True)
    ap.add_argument("--acta", type=str, default="[NÚMERO]")
    ap.add_argument("--sesion", type=str, required=True, help="Presencial/Mixta/Virtual")
    ap.add_argument("--fecha", type=str, default=dt.date.today().strftime("%d/%m/%Y"))
    ap.add_argument("--inicio", type=str, default="")
    ap.add_argument("--cierre", type=str, default="No consta")
    ap.add_argument("--lugar", type=str, required=True)

    args = ap.parse_args()
    outdir = Path(args.outdir); outdir.mkdir(parents=True, exist_ok=True)
    name = args.name

    # 1) Obtener transcripción diarizada (.spk.txt)
    if args.spk:
        spk_txt = Path(args.spk)
        assert spk_txt.is_file(), f"No existe {spk_txt}"
        final_spk = outdir / f"{name}.spk.txt"
        if final_spk.resolve() != spk_txt.resolve():
            final_spk.write_text(spk_txt.read_text(encoding="utf-8"), encoding="utf-8")
    else:
        wav = Path(args.wav); assert wav.is_file(), f"No existe WAV {wav}"
        # FFmpeg a 16k mono
        print(">> Convirtiendo a 16 kHz mono…")
        ff = which_ffmpeg()
        if not ff: raise RuntimeError("ffmpeg no encontrado en PATH")
        wav16 = outdir / f"{name}_16k.wav"
        rc = run_cmd([ff, "-y", "-i", str(wav), "-ac", "1", "-ar", "16000", "-c:a", "pcm_s16le", str(wav16)])
        if rc != 0 or not wav16.exists():
            raise RuntimeError("Falló ffmpeg")

        # Whisper (faster-whisper)
        print(">> Transcribiendo con Faster-Whisper…")
        os.environ["KMP_DUPLICATE_LIB_OK"] = "TRUE"
        rc = run_cmd([sys.executable, "-m", "scripts.transcribe_whisper_local",
                      "--file", str(wav16), "--model", "medium", "--lang", "es"])
        if rc != 0: raise RuntimeError("Falló transcripción")
        whisper_json = outdir / f"{name}_16k.whisper.json"
        if not whisper_json.exists():
            alt = wav16.with_suffix(".whisper.json")
            whisper_json = alt if alt.exists() else whisper_json
        if not whisper_json.exists():
            alt2 = Path("scripts") / f"{wav16.stem}.whisper.json"
            if alt2.exists(): whisper_json = alt2
        if not whisper_json.exists():
            raise RuntimeError("No se encontró JSON de Whisper generado.")

        # Diarización + merge (SpeechBrain)
        print(f">> Diarizando y fusionando (speakers={args.speakers})…")
        rc = run_cmd([sys.executable, "-m", "scripts.diarize_speechbrain_merge",
                      "--wav", str(wav16), "--json", str(whisper_json), "--speakers", str(args.speakers)])
        if rc != 0: raise RuntimeError("Falló diarización/fusión")

        spk_txt = Path(str(whisper_json).replace(".json", ".spk.txt"))
        if not spk_txt.exists():
            alt = Path("scripts") / spk_txt.name
            if alt.exists(): spk_txt = alt
        if not spk_txt.exists():
            raise RuntimeError("No se encontró .spk.txt de salida")
        final_spk = outdir / f"{name}.spk.txt"
        final_spk.write_text(spk_txt.read_text(encoding="utf-8"), encoding="utf-8")

    # 2) Generar ACTA con IA
    if not args.api_key:
        raise RuntimeError("Falta DEEPSEEK_API_KEY (--api-key o variable de entorno)")

    conv_text = read_text_or_docx(Path(args.convocatoria))
    ejemplo_text = read_text_or_docx(Path(args.acta_ejemplo)) if args.acta_ejemplo else ""
    prompt_base = Path(args.prompt).read_text(encoding="utf-8", errors="replace")
    transcript_text = (outdir / f"{name}.spk.txt").read_text(encoding="utf-8", errors="replace")

    system = (
        "Eres un redactor experto de actas universitarias. "
        "Redacta en tercera persona, sobria y clara. Sigue el estilo del documento de ejemplo si está disponible. "
        "No inventes datos: deja [DATO NO CONSTA] cuando no esté en convocatoria ni transcripción. "
        "Entrega: Acta en Markdown + Resumen ejecutivo + tabla de decisiones/votaciones + compromisos.\n\n"
        "Reglas de intervenciones:\n"
        "- La transcripción usa etiquetas 'Persona 1', 'Persona 2', etc.\n"
        "- En 'Desarrollo de la sesión' presenta intervenciones así: *Persona 1:* ...  *Persona 2:* ...\n"
        "- Si se identifica un nombre/rol explícito en la transcripción, reemplaza 'Persona X' por ese nombre/rol; "
        "si no, conserva 'Persona X' o 'Interviniente [X]'.\n\n"
        "--- PROMPT BASE ---\n" + prompt_base
    )

    meta = {
        "COMITÉ": args.comite, "ACTA": args.acta or "[NÚMERO]", "SESIÓN": args.sesion, "FECHA": args.fecha,
        "HORA INICIO": args.inicio, "HORA CIERRE": args.cierre or "No consta", "LUGAR": args.lugar
    }
    user = (
        "## Metadatos\n" + "\n".join(f"- {k}: {v}" for k,v in meta.items()) +
        "\n\n## Convocatoria (texto)\n" + conv_text[:120000] +
        "\n\n## Documento de Ejemplo (texto)\n" + (ejemplo_text[:120000] if ejemplo_text else "[No provisto]") +
        "\n\n## Transcripción diarizada\n" + transcript_text[:350000] +
        "\n\n### Tarea\nElabora el acta completa (Markdown) con el orden del día de la convocatoria, "
        "ajustando el desarrollo a lo indicado por la transcripción. Incluye votaciones/decisiones y compromisos."
    )

    print(">> Generando acta con DeepSeek…")
    md = call_deepseek_chat(
        messages=[{"role":"system","content":system},{"role":"user","content":user}],
        api_key=args.api_key
    )
    md_path = outdir / f"{name}.acta.md"; md_path.write_text(md, encoding="utf-8")
    docx_path = outdir / f"{name}.acta.docx"; md_to_docx(md, docx_path)

    print(f"\n✔ Hecho.\n- Transcripción: {outdir / (name + '.spk.txt')}\n- Acta MD: {md_path}\n- Acta DOCX: {docx_path}")

if __name__ == "__main__":
    main()
